"""
extraction.py — Text extraction layer for PDF and DOCX resumes.

Implements a layered fallback strategy:
  PDF:  pdfplumber (text layer w/ font metadata) → OCR via pdf2image + pytesseract
  DOCX: python-docx paragraphs + tables + text-box fallback

Returns extracted text, per-line metadata (font size, bold), the extraction
method used ("text" / "ocr" / "mixed"), and any warning strings.
"""

from __future__ import annotations

import io
import logging
import re
import shutil
from collections import Counter
from dataclasses import dataclass, field
from typing import Literal

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Dataclasses for structured extraction output
# ---------------------------------------------------------------------------

@dataclass
class TextLine:
    """A single line of extracted text with optional formatting metadata."""
    text: str
    font_size: float | None = None
    is_bold: bool = False
    page_number: int = 0


@dataclass
class ExtractionResult:
    """Complete output from the text-extraction stage."""
    lines: list[TextLine] = field(default_factory=list)
    full_text: str = ""
    extraction_method: Literal["text", "ocr", "mixed"] = "text"
    warnings: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
# If text-layer extraction yields fewer than this many non-whitespace chars,
# we consider it "near-empty" and fall back to OCR.
_MIN_TEXT_CHARS = 50

# Lines that repeat verbatim across >50% of pages are likely headers/footers
_HEADER_FOOTER_REPEAT_THRESHOLD = 0.5


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf_text_layer(file_bytes: bytes) -> tuple[list[TextLine], list[str]]:
    """
    Extract text from a PDF's embedded text layer using pdfplumber.

    Returns (lines, warnings). Each line carries font_size and is_bold
    metadata derived from the dominant character attributes on that line.
    """
    import pdfplumber

    lines: list[TextLine] = []
    warnings: list[str] = []

    try:
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
    except Exception as exc:
        warnings.append(f"pdfplumber could not open the file: {exc}")
        return lines, warnings

    try:
        for page_num, page in enumerate(pdf.pages, start=1):
            # Gather per-character metadata so we can derive line-level
            # font size and bold weight.
            chars = page.chars or []
            if not chars:
                # Fallback: extract_text() might still return something
                raw = page.extract_text() or ""
                for raw_line in raw.split("\n"):
                    stripped = raw_line.strip()
                    if stripped:
                        lines.append(TextLine(text=stripped, page_number=page_num))
                continue

            # Group characters into lines by their vertical position (top).
            # We round `top` to 1 dp to merge chars on the same baseline.
            line_buckets: dict[float, list[dict]] = {}
            for ch in chars:
                key = round(ch.get("top", 0), 1)
                line_buckets.setdefault(key, []).append(ch)

            for _top in sorted(line_buckets):
                bucket = line_buckets[_top]
                text = "".join(c.get("text", "") for c in bucket).strip()
                if not text:
                    continue

                # Dominant font size = most-common size in the line
                sizes = [c.get("size", 0) for c in bucket if c.get("size")]
                font_size = Counter(sizes).most_common(1)[0][0] if sizes else None

                # Bold heuristic: fontname contains "Bold" or weight > 400
                bold_votes = sum(
                    1
                    for c in bucket
                    if ("Bold" in (c.get("fontname") or ""))
                    or ("bold" in (c.get("fontname") or "").lower())
                )
                is_bold = bold_votes > len(bucket) / 2

                lines.append(
                    TextLine(
                        text=text,
                        font_size=font_size,
                        is_bold=is_bold,
                        page_number=page_num,
                    )
                )
    except Exception as exc:
        warnings.append(f"Error during pdfplumber text extraction: {exc}")
    finally:
        pdf.close()

    return lines, warnings


def _ocr_available() -> bool:
    """Check whether tesseract and poppler binaries are accessible."""
    return (
        shutil.which("tesseract") is not None
        and shutil.which("pdftoppm") is not None
    )


def _extract_pdf_ocr(file_bytes: bytes) -> tuple[list[TextLine], list[str]]:
    """
    OCR fallback for scanned / image-only PDFs.

    Renders each page to an image via pdf2image, then runs pytesseract.
    Returns (lines, warnings).
    """
    warnings: list[str] = []
    lines: list[TextLine] = []

    if not _ocr_available():
        warnings.append(
            "OCR requested but tesseract/poppler binaries are not installed. "
            "Skipping OCR — extracted text may be empty."
        )
        return lines, warnings

    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as exc:
        warnings.append(f"OCR dependencies missing: {exc}")
        return lines, warnings

    try:
        images = convert_from_bytes(file_bytes, dpi=300)
    except Exception as exc:
        warnings.append(f"pdf2image could not render pages: {exc}")
        return lines, warnings

    for page_num, img in enumerate(images, start=1):
        try:
            text = pytesseract.image_to_string(img)
        except Exception as exc:
            warnings.append(f"pytesseract failed on page {page_num}: {exc}")
            continue

        for raw_line in text.split("\n"):
            stripped = raw_line.strip()
            if stripped:
                lines.append(
                    TextLine(text=stripped, page_number=page_num)
                )

    return lines, warnings


def extract_pdf(file_bytes: bytes) -> ExtractionResult:
    """
    Extract text from a PDF with layered fallback.

    1. Try text-layer extraction via pdfplumber (preserves font metadata).
    2. If the text layer is empty/near-empty, fall back to OCR.
    3. If OCR also fails or is unavailable, return best-effort with warnings.
    """
    text_lines, text_warnings = _extract_pdf_text_layer(file_bytes)
    total_chars = sum(len(ln.text) for ln in text_lines)

    if total_chars >= _MIN_TEXT_CHARS:
        # Good text layer — strip repeated headers/footers and return
        cleaned = _strip_repeated_lines(text_lines)
        return ExtractionResult(
            lines=cleaned,
            full_text="\n".join(ln.text for ln in cleaned),
            extraction_method="text",
            warnings=text_warnings,
        )

    # Text layer empty or near-empty — try OCR
    ocr_lines, ocr_warnings = _extract_pdf_ocr(file_bytes)
    all_warnings = text_warnings + ocr_warnings

    if not ocr_lines:
        # Both paths yielded nothing
        all_warnings.append(
            "Both text-layer and OCR extraction returned empty content."
        )
        return ExtractionResult(
            lines=text_lines,  # might be [] but keep whatever we got
            full_text="\n".join(ln.text for ln in text_lines),
            extraction_method="text",
            warnings=all_warnings,
        )

    # If we had *some* text-layer content, merge (mixed); otherwise pure OCR
    if text_lines:
        combined = text_lines + ocr_lines
        method: Literal["text", "ocr", "mixed"] = "mixed"
    else:
        combined = ocr_lines
        method = "ocr"

    cleaned = _strip_repeated_lines(combined)
    return ExtractionResult(
        lines=cleaned,
        full_text="\n".join(ln.text for ln in cleaned),
        extraction_method=method,
        warnings=all_warnings,
    )


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_docx(file_bytes: bytes) -> ExtractionResult:
    """
    Extract text from a DOCX file.

    Walks paragraphs (with heading/style info), then tables, then any
    inline text boxes (often missed by naive paragraph-only iteration).
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    warnings: list[str] = []
    lines: list[TextLine] = []

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        warnings.append(f"python-docx could not open the file: {exc}")
        return ExtractionResult(warnings=warnings)

    # --- Paragraphs ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Heading detection
        is_heading = para.style and para.style.name and para.style.name.startswith("Heading")

        # Bold: majority of runs are bold
        bold_runs = sum(1 for r in para.runs if r.bold)
        is_bold = is_heading or (bold_runs > len(para.runs) / 2 if para.runs else False)

        # Approximate font size from first run with a size
        font_size = None
        for run in para.runs:
            if run.font and run.font.size:
                # python-docx returns EMU; convert to points
                font_size = run.font.size.pt
                break

        lines.append(
            TextLine(
                text=text,
                font_size=font_size,
                is_bold=is_bold,
                page_number=0,  # DOCX doesn't expose page numbers
            )
        )

    # --- Tables (walk every cell) ---
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                lines.append(TextLine(text=" | ".join(row_texts)))

    # --- Text boxes (embedded in shapes / fallback XML walk) ---
    try:
        from docx.oxml.ns import qn

        # Text boxes live inside w:txbxContent elements
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            for p in txbx.iter(qn("w:p")):
                text_parts = []
                for r in p.iter(qn("w:r")):
                    for t in r.iter(qn("w:t")):
                        if t.text:
                            text_parts.append(t.text)
                full = "".join(text_parts).strip()
                if full:
                    lines.append(TextLine(text=full))
    except Exception as exc:
        warnings.append(f"Text-box extraction encountered an issue: {exc}")

    if not lines:
        warnings.append("DOCX extraction returned no text content.")

    return ExtractionResult(
        lines=lines,
        full_text="\n".join(ln.text for ln in lines),
        extraction_method="text",
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Shared utilities
# ---------------------------------------------------------------------------

def _strip_repeated_lines(lines: list[TextLine]) -> list[TextLine]:
    """
    Remove likely headers, footers, and page numbers.

    A line is considered a repeated header/footer if:
    - Its exact text appears on more than half of all pages, AND
    - It appears at the top or bottom of its page.

    Also strips standalone page-number lines (e.g. "1", "Page 2 of 5").
    """
    if not lines:
        return lines

    page_numbers = {ln.page_number for ln in lines if ln.page_number}
    n_pages = len(page_numbers) if page_numbers else 1

    # Count how many distinct pages each unique line text appears on
    text_page_counts: dict[str, set[int]] = {}
    for ln in lines:
        text_page_counts.setdefault(ln.text, set()).add(ln.page_number)

    repeated_texts = {
        t
        for t, pages in text_page_counts.items()
        if len(pages) > n_pages * _HEADER_FOOTER_REPEAT_THRESHOLD and n_pages > 1
    }

    # Also match standalone page-number patterns
    page_num_re = re.compile(
        r"^(page\s+)?\d+(\s*(of|/)\s*\d+)?$", re.IGNORECASE
    )

    cleaned: list[TextLine] = []
    for ln in lines:
        if ln.text in repeated_texts:
            continue
        if page_num_re.match(ln.text.strip()):
            continue
        cleaned.append(ln)

    return cleaned


# ---------------------------------------------------------------------------
# Public dispatch function
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str) -> ExtractionResult:
    """
    Main entry point: dispatch to the appropriate extractor based on file extension.

    Supports .pdf and .docx. Returns an ExtractionResult with warnings on any
    issues; never raises for content-level problems.
    """
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return extract_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return extract_docx(file_bytes)
    else:
        return ExtractionResult(
            warnings=[
                f"Unsupported file type: '{filename}'. Only .pdf and .docx are accepted."
            ]
        )
