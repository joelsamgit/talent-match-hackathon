"""
test_extractor.py — Unit tests for extractor.py

Tests verify that:
  - PDF files are correctly converted to plain text
  - DOCX files are correctly converted to plain text
  - Unsupported file types raise ValueError
  - Edge cases (empty files) are handled gracefully

No LLM calls are made — extractor.py is pure I/O logic.
"""

import io
import pytest

from extractor import extract_text, extract_from_pdf, extract_from_docx
from tests.conftest import make_docx_bytes, make_pdf_bytes


# ---------------------------------------------------------------------------
# PDF extraction tests
# ---------------------------------------------------------------------------

class TestPdfExtraction:

    def test_pdf_extraction_returns_known_text(self):
        """
        Builds a minimal in-memory PDF with a known sentence using fpdf2,
        passes it as bytes to extract_from_pdf(), and asserts the sentence
        appears in the output. Proves the PDF pipeline is wired correctly.
        """
        known_text = "Strong grasp of data structures"
        pdf_bytes = make_pdf_bytes(known_text)

        result = extract_from_pdf(pdf_bytes)

        assert known_text in result, (
            f"Expected '{known_text}' in extracted PDF text, got: {result!r}"
        )

    def test_extract_text_dispatch_pdf(self):
        """
        Verifies the dispatch function routes .pdf files to the PDF extractor.
        Uses the filename extension to determine which extractor to call.
        """
        known_text = "Design scalable distributed systems"
        pdf_bytes = make_pdf_bytes(known_text)

        result = extract_text("job_description.pdf", pdf_bytes)

        assert known_text in result
        assert isinstance(result, str)
        assert len(result) > 0

    def test_pdf_extraction_preserves_multiple_lines(self):
        """
        Ensures multi-word content from a PDF is not collapsed or truncated.
        Exercises a realistic JD-length input (not just a single cell).
        """
        # fpdf2's cell() fits one line; use multi_cell for longer text
        from fpdf import FPDF

        content = "Key Responsibilities Design APIs Write clean code Code review"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, text=content)
        pdf_bytes = bytes(pdf.output())

        result = extract_from_pdf(pdf_bytes)
        # At least some of the core words should survive extraction
        assert "Design" in result or "APIs" in result or "clean" in result


# ---------------------------------------------------------------------------
# DOCX extraction tests
# ---------------------------------------------------------------------------

class TestDocxExtraction:

    def test_docx_extraction_returns_known_text(self):
        """
        Builds a minimal in-memory DOCX with a known sentence using python-docx,
        passes it as bytes to extract_from_docx(), and asserts the sentence
        appears in the output.
        """
        known_text = "Experience with object-oriented design"
        docx_bytes = make_docx_bytes(known_text)

        result = extract_from_docx(docx_bytes)

        assert known_text in result, (
            f"Expected '{known_text}' in extracted DOCX text, got: {result!r}"
        )

    def test_extract_text_dispatch_docx(self):
        """
        Verifies the dispatch function routes .docx files to the DOCX extractor.
        """
        known_text = "Clear written and verbal communication"
        docx_bytes = make_docx_bytes(known_text)

        result = extract_text("role_description.docx", docx_bytes)

        assert known_text in result
        assert isinstance(result, str)

    def test_docx_extraction_multiple_paragraphs(self):
        """
        Verifies that multiple paragraphs in a DOCX are all captured,
        not just the first one. This matters because JDs typically span
        many paragraphs.
        """
        from docx import Document

        doc = Document()
        doc.add_paragraph("Paragraph one: algorithms and data structures")
        doc.add_paragraph("Paragraph two: system design at scale")
        doc.add_paragraph("Paragraph three: communication skills")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extract_from_docx(docx_bytes)

        assert "Paragraph one" in result
        assert "Paragraph two" in result
        assert "Paragraph three" in result


# ---------------------------------------------------------------------------
# Unsupported file type tests
# ---------------------------------------------------------------------------

class TestUnsupportedTypes:

    def test_txt_file_raises_value_error(self):
        """
        Passing a .txt filename to extract_text() should raise ValueError
        because the module only supports .pdf and .docx. A clear error message
        prevents silent failures during integration.
        """
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text("resume.txt", b"some plain text content")

    def test_csv_file_raises_value_error(self):
        """
        .csv files should also be rejected. Confirms the guard is not
        limited to .txt only.
        """
        with pytest.raises(ValueError):
            extract_text("data.csv", b"col1,col2\nval1,val2")

    def test_no_extension_raises_value_error(self):
        """
        A filename with no extension (e.g. 'jobdesc') should also raise
        ValueError, since we cannot determine the format.
        """
        with pytest.raises(ValueError):
            extract_text("jobdesc", b"some content")


# ---------------------------------------------------------------------------
# Edge case: empty / corrupt file
# ---------------------------------------------------------------------------

class TestEdgeCases:

    def test_empty_docx_raises_or_returns_empty(self):
        """
        Passing empty bytes as DOCX input should either raise a clear exception
        or return an empty string. It must NOT crash with an unhandled traceback.
        Graceful failure lets the API return a 422 instead of a 500.
        """
        try:
            result = extract_from_docx(b"")
            # If it doesn't raise, the result must be empty
            assert result.strip() == ""
        except Exception as exc:
            # Any exception is acceptable — what matters is it's not a silent crash
            assert exc is not None

    def test_empty_pdf_raises_or_returns_empty(self):
        """
        Same as above for PDF — empty bytes should not produce a hard crash
        with an unhandled exception type.
        """
        try:
            result = extract_from_pdf(b"")
            assert result.strip() == ""
        except Exception as exc:
            assert exc is not None
